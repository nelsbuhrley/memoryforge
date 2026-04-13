"""Tests for material parsers."""

from pathlib import Path

from memoryforge.parser.text_parser import parse_text, parse_markdown
from memoryforge.parser.pdf_parser import parse_pdf
from memoryforge.parser.docx_parser import parse_docx


class TestTextParser:
    def test_parse_plain_text(self, tmp_path: Path):
        content = "Chapter 1: Cells\n\nCells are the basic unit of life.\n\nChapter 2: DNA\n\nDNA stores genetic information."
        f = tmp_path / "notes.txt"
        f.write_text(content)
        result = parse_text(f)
        assert result.text == content
        assert result.page_count is None

    def test_parse_markdown(self, tmp_path: Path):
        f = tmp_path / "notes.md"
        f.write_text("# Chapter 1: Cells\n\nCells are the basic unit of life.\n\n# Chapter 2: DNA\n\nDNA stores genetic information.")
        result = parse_markdown(f)
        assert len(result.sections) == 2
        assert result.sections[0]["heading"] == "Chapter 1: Cells"

    def test_empty_file(self, tmp_path: Path):
        f = tmp_path / "empty.txt"
        f.write_text("")
        result = parse_text(f)
        assert result.text == ""

    def test_markdown_content_before_first_heading_is_excluded_from_sections(self, tmp_path: Path):
        f = tmp_path / "notes.md"
        f.write_text("Preamble text\n\n# Chapter 1: Cells\n\nCells are the basic unit of life.")
        result = parse_markdown(f)
        assert len(result.sections) == 1
        assert result.sections[0]["heading"] == "Chapter 1: Cells"


class TestPdfParser:
    def test_parse_pdf_returns_text(self, tmp_path: Path):
        """Integration test — requires pymupdf to create a test PDF."""
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Chapter 1: Introduction\n\nThis is the introduction.")
        pdf_path = tmp_path / "test.pdf"
        doc.save(str(pdf_path))
        doc.close()

        result = parse_pdf(pdf_path)
        assert result.page_count == 1
        assert "Introduction" in result.text

    def test_parse_pdf_structure(self, tmp_path: Path):
        import fitz
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content")
        pdf_path = tmp_path / "multi.pdf"
        doc.save(str(pdf_path))
        doc.close()

        result = parse_pdf(pdf_path)
        assert result.page_count == 3


class TestDocxParser:
    def test_parse_docx(self, tmp_path: Path):
        from docx import Document
        doc = Document()
        doc.add_heading("Chapter 1: Cells", level=1)
        doc.add_paragraph("Cells are the basic unit of life.")
        doc.add_heading("Chapter 2: DNA", level=1)
        doc.add_paragraph("DNA stores genetic information.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = parse_docx(docx_path)
        assert "Cells" in result.text
        assert len(result.sections) == 2

    def test_parse_docx_plain_paragraph_not_in_sections(self, tmp_path: Path):
        """Plain (non-heading) paragraphs appear in text but not in sections."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Plain paragraph")
        docx_path = tmp_path / "plain.docx"
        doc.save(str(docx_path))

        result = parse_docx(docx_path)
        assert "Plain paragraph" in result.text
        assert result.sections == []
